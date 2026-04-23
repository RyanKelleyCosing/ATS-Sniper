#!/usr/bin/env python3
"""Generate one or more tailored application packs from a JD or markdown manifest."""

from __future__ import annotations

import argparse
import csv
import json
from datetime import datetime
from dataclasses import dataclass
from pathlib import Path
import re
import sys
from typing import Any, Mapping

from docx import Document
from docx.oxml.ns import qn
from openai import OpenAI
from pypdf import PdfReader

from generate_tailored_resume import generate_tailored_resume_for_job
from job_scraper import fetch_job_description
from resume_tailor import CoverLetterGuidance, generate_cover_letter
from startup_discovery_scraper import run_jobspy_queries
from utils.application_batch import (
    ApplicationBatchJob,
    ApplicationBatchManifest,
    build_tailored_job_description,
    is_usable_job_description,
    parse_application_batch_manifest,
)
from utils.application_packages import package_application_artifacts
from utils.state import load_config, load_state

DEFAULT_PACKAGE_ROOT = Path("resumes to make manually") / "Application Packs"
DEFAULT_REVIEW_PACKAGE_ROOT = DEFAULT_PACKAGE_ROOT / "Review Queue"
DEFAULT_MATCH_SCORE = 80
DEFAULT_JOB_URL = "manual://pasted-job-description"
DEFAULT_APPLICATION_PACKAGE_MODEL = "gpt-5.4"
ATS_AUDIT_CHECK_KEYS = (
    "relevance",
    "must_have_keyword_gate",
    "uniqueness",
    "company_keyword_gate",
    "keyword_density",
    "keyword_role_spread",
    "impact_first_bullets",
    "action_context_result",
)
ATS_STANDARD_HEADINGS = frozenset(
    {
        "Skills",
        "Technical Skills",
        "TECHNICAL SKILLS",
        "Experience",
        "Work Experience",
        "WORK EXPERIENCE",
        "Professional Summary",
        "PROFESSIONAL SUMMARY",
        "Education",
        "EDUCATION",
        "Certifications",
        "CERTIFICATIONS",
        "Projects",
    }
)


@dataclass(frozen=True)
class ManualApplicationRequest:
    """Inputs for a manual application package generation run."""

    company: str
    role: str
    job_description: str
    job_url: str
    match_score: int
    package_root: Path
    template_hint: str | None = None
    preferred_accomplishment_ids: tuple[str, ...] = ()
    industry_focus: str = ""
    cover_letter_focus: str = ""
    tailoring_notes: tuple[str, ...] = ()
    include_supporting_artifacts: bool = False
    allow_must_have_keyword_override: bool = False


@dataclass(frozen=True)
class ApplicationGenerationContext:
    """Shared OpenAI context for a manual or batch application run."""

    client: OpenAI
    model: str


def parse_arguments() -> argparse.Namespace:
    """Parse command-line arguments for one-off or batch application generation."""
    parser = argparse.ArgumentParser(
        description=(
            "Generate tailored resumes and cover letters from a single job description "
            "or a markdown batch manifest."
        )
    )
    parser.add_argument("--batch-file", help="Markdown or YAML file describing one or more jobs")
    parser.add_argument("--company", help="Target company name")
    parser.add_argument("--role", help="Target role or title")
    parser.add_argument("--job-url", default=DEFAULT_JOB_URL, help="Optional job posting URL")
    parser.add_argument(
        "--match-score",
        type=int,
        default=DEFAULT_MATCH_SCORE,
        help="Manual fit score to pass into the tailored-resume generator",
    )
    parser.add_argument("--job-description", help="Job description text passed directly on the command line")
    parser.add_argument("--job-description-file", help="Path to a text file containing the full job description")
    parser.add_argument(
        "--package-root",
        default=str(DEFAULT_PACKAGE_ROOT),
        help="Where the final application pack should be written",
    )
    parser.add_argument("--template-hint", help="Optional resume template hint such as cloud, devops, infrastructure, or sre")
    parser.add_argument(
        "--preferred-accomplishments",
        help="Comma-separated accomplishment IDs that must be included when truthful",
    )
    parser.add_argument("--industry-focus", help="Industry framing to emphasize in the cover letter")
    parser.add_argument("--cover-letter-focus", help="Role or company nuance to emphasize in the cover letter")
    parser.add_argument(
        "--tailoring-note",
        action="append",
        default=[],
        help="Additional note to append to the JD before tailoring; repeat for multiple notes",
    )
    parser.add_argument(
        "--include-supporting-artifacts",
        action="store_true",
        help="Keep analysis, source markdown, apply shortcut, and review checklist in each package",
    )
    parser.add_argument(
        "--model",
        help="Override the OpenAI model for resume and cover-letter generation",
    )
    parser.add_argument(
        "--min-match-score",
        type=int,
        help="When used with --batch-file, only process jobs at or above this historical match score.",
    )
    parser.add_argument(
        "--max-match-score",
        type=int,
        help="When used with --batch-file, only process jobs at or below this historical match score.",
    )
    return parser.parse_args()


def prompt_for_job_description() -> str:
    """Read a pasted multi-line job description from standard input."""
    print(
        "Paste the full job description, then finish with Ctrl+Z and Enter on Windows "
        "or Ctrl+D on macOS/Linux."
    )
    return sys.stdin.read().strip()


def load_job_description_from_args(args: argparse.Namespace) -> str:
    """Resolve the one-off job description from CLI text, a file, or pasted stdin."""
    if args.job_description:
        return args.job_description.strip()
    if args.job_description_file:
        return Path(args.job_description_file).read_text(encoding="utf-8").strip()
    return prompt_for_job_description()


def parse_preferred_accomplishments(raw_value: str | None) -> tuple[str, ...]:
    """Parse a comma-separated accomplishment list from the command line."""
    if not raw_value:
        return ()
    return tuple(item.strip() for item in raw_value.split(',') if item.strip())


def parse_queue_ranks(raw_value: str | None) -> tuple[int, ...]:
    """Parse a comma-separated queue-rank list from the command line."""
    if not raw_value:
        return ()

    queue_ranks: list[int] = []
    for item in raw_value.split(','):
        normalized_item = str(item).strip()
        if not normalized_item:
            continue
        queue_ranks.append(int(normalized_item))
    return tuple(queue_ranks)


def _is_truthy_csv_value(value: Any) -> bool:
    """Return True when a CSV field is an affirmative-like value."""
    normalized = str(value or "").strip().casefold()
    return normalized in {"1", "true", "yes", "y"}


def _normalize_review_match_text(value: str) -> str:
    """Normalize a company or title string for loose review-row matching."""
    without_parentheticals = re.sub(r"\s*\([^)]*\)", " ", str(value))
    return " ".join(without_parentheticals.casefold().split())


def infer_jobspy_site_name(review_row: Mapping[str, Any]) -> str:
    """Infer a JobSpy-supported board site from the review CSV row."""
    source_board = str(review_row.get("Source Board", "")).strip().casefold()
    if source_board in {"indeed", "google", "zip_recruiter"}:
        return source_board

    job_url = str(review_row.get("URL", "")).strip().casefold()
    if "indeed.com" in job_url:
        return "indeed"
    if "google" in job_url:
        return "google"
    if "ziprecruiter.com" in job_url:
        return "zip_recruiter"
    return ""


def fetch_review_description_from_jobspy(review_row: Mapping[str, Any]) -> str:
    """Fetch a blocked board job description through JobSpy using the review row metadata."""
    site_name = infer_jobspy_site_name(review_row)
    if not site_name:
        return ""

    company = str(review_row.get("Company", "")).strip()
    role = str(review_row.get("Title", "")).strip()
    job_url = str(review_row.get("URL", "")).strip()
    location = str(review_row.get("Location", "United States")).strip() or "United States"
    is_remote = "remote" in location.casefold()
    search_terms = tuple(
        dict.fromkeys(
            term
            for term in (
                role,
                re.sub(r"\s*\([^)]*\)", "", role).strip(),
            )
            if term
        )
    )

    query_requests = [
        {
            "search_kwargs": {
                "site_name": [site_name],
                "search_term": search_term,
                "location": "United States" if is_remote else location,
                "hours_old": 168,
                "country_indeed": "USA",
                "is_remote": is_remote,
                "results_wanted": 10,
            }
        }
        for search_term in search_terms
    ]
    query_results = run_jobspy_queries(query_requests)
    normalized_company = _normalize_review_match_text(company)
    normalized_role = _normalize_review_match_text(role)
    fallback_description = ""

    for query_result in query_results:
        for record in query_result.get("records", []):
            if not isinstance(record, Mapping):
                continue
            description = str(record.get("description") or "").strip()
            if not description:
                continue

            record_url = str(record.get("job_url") or record.get("url") or "").strip()
            record_company = _normalize_review_match_text(str(record.get("company") or ""))
            record_title = _normalize_review_match_text(str(record.get("title") or ""))
            if job_url and record_url == job_url:
                return description
            if normalized_company and normalized_role and record_company == normalized_company and record_title == normalized_role:
                return description
            if normalized_company and record_company == normalized_company and not fallback_description:
                fallback_description = description

    return fallback_description


def load_review_csv_rows(review_csv_path: Path) -> list[dict[str, str]]:
    """Load the review-queue CSV exported by the hot-job processor."""
    with review_csv_path.open("r", newline="", encoding="utf-8") as file_handle:
        return [dict(row) for row in csv.DictReader(file_handle)]


def filter_review_csv_rows(
    rows: list[dict[str, str]],
    *,
    queue_ranks: tuple[int, ...] = (),
    min_match_score: int | None = None,
    max_match_score: int | None = None,
    actionable_only: bool = True,
    limit: int | None = None,
) -> list[dict[str, str]]:
    """Filter and sort review CSV rows for package generation."""
    selected_queue_ranks = set(queue_ranks)
    filtered_rows: list[dict[str, str]] = []

    for row in rows:
        queue_rank_text = str(row.get("Queue Rank", "")).strip()
        match_score_text = str(row.get("Match Score", "")).strip()
        queue_rank = int(queue_rank_text) if queue_rank_text else None
        match_score = int(match_score_text) if match_score_text else DEFAULT_MATCH_SCORE

        if selected_queue_ranks and queue_rank not in selected_queue_ranks:
            continue
        if actionable_only and not _is_truthy_csv_value(row.get("Actionable Review", "")):
            continue
        if min_match_score is not None and match_score < min_match_score:
            continue
        if max_match_score is not None and match_score > max_match_score:
            continue

        filtered_rows.append(row)

    filtered_rows.sort(
        key=lambda row: (
            int(str(row.get("Queue Rank", "0")).strip() or 0),
            str(row.get("Company", "")).strip(),
            str(row.get("Title", "")).strip(),
        )
    )
    if limit is not None:
        return filtered_rows[:limit]
    return filtered_rows


def resolve_review_job_description(
    review_row: Mapping[str, Any],
    *,
    state: Mapping[str, Any],
) -> str:
    """Resolve job text for a review-queue row from state first, then the live URL."""
    review_row_description = str(
        review_row.get("Job Description Snapshot") or review_row.get("Job Description") or ""
    ).strip()
    if is_usable_job_description(review_row_description):
        return review_row_description

    job_url = str(review_row.get("URL", "")).strip()
    job_record = state.get("jobs", {}).get(job_url, {}) if isinstance(state.get("jobs", {}), Mapping) else {}

    state_description = str(
        job_record.get("job_description") or job_record.get("description") or ""
    ).strip() if isinstance(job_record, Mapping) else ""
    if is_usable_job_description(state_description):
        return state_description

    fetched_description = fetch_remote_job_description(job_url)
    if is_usable_job_description(fetched_description):
        return fetched_description

    jobspy_description = fetch_review_description_from_jobspy(review_row)
    if is_usable_job_description(jobspy_description):
        return jobspy_description

    if review_row_description:
        return review_row_description
    if state_description:
        return state_description
    if fetched_description:
        return fetched_description
    if jobspy_description:
        return jobspy_description

    raise ValueError(f"No usable job description was available for review row URL: {job_url or 'missing URL'}")


def build_request_from_review_csv_row(
    review_row: Mapping[str, Any],
    *,
    state: Mapping[str, Any],
    package_root: Path,
    include_supporting_artifacts: bool,
) -> ManualApplicationRequest:
    """Build a package-generation request from one review-queue CSV row."""
    company = str(review_row.get("Company", "")).strip()
    role = str(review_row.get("Title", "")).strip()
    if not company or not role:
        raise ValueError("Review CSV rows must include Company and Title.")

    job_url = str(review_row.get("URL", DEFAULT_JOB_URL)).strip() or DEFAULT_JOB_URL
    match_score_text = str(review_row.get("Match Score", "")).strip()
    match_score = int(match_score_text) if match_score_text else DEFAULT_MATCH_SCORE

    return ManualApplicationRequest(
        company=company,
        role=role,
        job_description=resolve_review_job_description(review_row, state=state),
        job_url=job_url,
        match_score=match_score,
        package_root=package_root,
        include_supporting_artifacts=include_supporting_artifacts,
        allow_must_have_keyword_override=True,
    )


def build_manifest_from_review_csv_rows(
    review_rows: list[dict[str, str]],
    *,
    package_root: Path,
    include_supporting_artifacts: bool,
) -> ApplicationBatchManifest:
    """Build a lightweight manifest-like wrapper for review CSV batch summaries."""
    jobs = tuple(
        ApplicationBatchJob(
            company=str(row.get("Company", "Unknown")).strip() or "Unknown",
            role=str(row.get("Title", "Unknown")).strip() or "Unknown",
            job_url=str(row.get("URL", DEFAULT_JOB_URL)).strip() or DEFAULT_JOB_URL,
            match_score=int(str(row.get("Match Score", DEFAULT_MATCH_SCORE)).strip() or DEFAULT_MATCH_SCORE),
            template_hint=None,
            preferred_accomplishment_ids=(),
            industry_focus="",
            cover_letter_focus="",
            job_description="",
            fallback_job_description="",
            tailoring_notes=(),
            package_root=None,
            include_supporting_artifacts=include_supporting_artifacts,
        )
        for row in review_rows
    )
    return ApplicationBatchManifest(
        package_root=package_root,
        include_supporting_artifacts=include_supporting_artifacts,
        jobs=jobs,
    )


def build_generation_context(model_override: str | None = None) -> ApplicationGenerationContext:
    """Create the shared OpenAI client and model selection for application generation."""
    config = load_config()
    settings = config.get("settings", {})
    model = (
        (model_override or "").strip()
        or str(settings.get("application_package_model", "")).strip()
        or DEFAULT_APPLICATION_PACKAGE_MODEL
    )
    return ApplicationGenerationContext(
        client=OpenAI(api_key=config.get("openai_key")),
        model=model,
    )


def resolve_package_root_override(raw_package_root: str | None) -> Path | None:
    """Return an explicit package-root override, ignoring the CLI default value."""
    if not raw_package_root:
        return None

    candidate = Path(raw_package_root)
    if candidate == DEFAULT_PACKAGE_ROOT:
        return None
    return candidate


def build_request(args: argparse.Namespace) -> ManualApplicationRequest:
    """Build a validated one-off request object from CLI arguments."""
    company = str(args.company or "").strip()
    role = str(args.role or "").strip()
    if not company or not role:
        raise ValueError("Single-job runs require both --company and --role.")

    job_description = build_tailored_job_description(
        load_job_description_from_args(args),
        tuple(str(note).strip() for note in args.tailoring_note if str(note).strip()),
        str(args.industry_focus or "").strip(),
    )
    if not job_description:
        raise ValueError("Job description text is required.")

    return ManualApplicationRequest(
        company=company,
        role=role,
        job_description=job_description,
        job_url=(args.job_url or DEFAULT_JOB_URL).strip(),
        match_score=args.match_score,
        package_root=Path(args.package_root),
        template_hint=(str(args.template_hint or "").strip() or None),
        preferred_accomplishment_ids=parse_preferred_accomplishments(args.preferred_accomplishments),
        industry_focus=str(args.industry_focus or "").strip(),
        cover_letter_focus=str(args.cover_letter_focus or "").strip(),
        tailoring_notes=tuple(str(note).strip() for note in args.tailoring_note if str(note).strip()),
        include_supporting_artifacts=bool(args.include_supporting_artifacts),
    )


def fetch_remote_job_description(job_url: str) -> str:
    """Fetch the live job description text from a supported public URL."""
    if not job_url.lower().startswith(("http://", "https://")):
        return ""

    job_data = fetch_job_description(job_url) or {}
    return str(job_data.get("description", "")).strip()


def resolve_batch_job_description(job: ApplicationBatchJob) -> str:
    """Resolve the best available description text for a batch-manifest job."""
    sections: list[str] = []
    fetched_description = fetch_remote_job_description(job.job_url)
    if is_usable_job_description(fetched_description):
        sections.append(fetched_description)
    if job.job_description:
        sections.append(job.job_description)
    if not sections and job.fallback_job_description:
        sections.append(job.fallback_job_description)
    if not sections:
        raise ValueError(f"No usable job description was available for {job.company} | {job.role}.")

    return build_tailored_job_description(
        "\n\n".join(section.strip() for section in sections if section.strip()),
        job.tailoring_notes,
        job.industry_focus,
    )


def build_request_from_batch_job(
    job: ApplicationBatchJob,
    manifest: ApplicationBatchManifest,
    package_root_override: Path | None = None,
    include_supporting_artifacts_override: bool | None = None,
) -> ManualApplicationRequest:
    """Build a concrete generation request from a batch-manifest job entry."""
    return ManualApplicationRequest(
        company=job.company,
        role=job.role,
        job_description=resolve_batch_job_description(job),
        job_url=job.job_url,
        match_score=job.match_score,
        package_root=package_root_override or job.package_root or manifest.package_root,
        template_hint=job.template_hint,
        preferred_accomplishment_ids=job.preferred_accomplishment_ids,
        industry_focus=job.industry_focus,
        cover_letter_focus=job.cover_letter_focus,
        tailoring_notes=job.tailoring_notes,
        include_supporting_artifacts=(
            include_supporting_artifacts_override
            if include_supporting_artifacts_override is not None
            else (
                manifest.include_supporting_artifacts
                if job.include_supporting_artifacts is None
                else job.include_supporting_artifacts
            )
        ),
    )


def filter_batch_jobs_by_match_score(
    jobs: tuple[ApplicationBatchJob, ...],
    *,
    min_match_score: int | None = None,
    max_match_score: int | None = None,
) -> tuple[ApplicationBatchJob, ...]:
    """Return only jobs whose manifest match scores fall within the requested range."""
    filtered_jobs: list[ApplicationBatchJob] = []
    for job in jobs:
        if min_match_score is not None and job.match_score < min_match_score:
            continue
        if max_match_score is not None and job.match_score > max_match_score:
            continue
        filtered_jobs.append(job)
    return tuple(filtered_jobs)


def build_ats_audit_summary(result: dict[str, object]) -> dict[str, object]:
    """Collect the ATS checks and artifact-level DOCX/PDF audit for a generated package."""
    checks: dict[str, dict[str, object]] = {}
    for key in ATS_AUDIT_CHECK_KEYS:
        raw_check = result.get(key)
        if not isinstance(raw_check, dict):
            continue
        check_summary: dict[str, object] = {
            "status": raw_check.get("status"),
            "message": raw_check.get("message"),
        }
        for metric_key in (
            "coverage_ratio",
            "role_spread_ratio",
            "impact_ratio",
            "acr_ratio",
            "mentioned_terms",
            "target_phrases",
            "matched_phrases",
            "missing_phrases",
            "missing_keywords",
            "matched_count",
            "required_matches",
        ):
            metric_value = raw_check.get(metric_key)
            if metric_value is not None:
                check_summary[metric_key] = metric_value
        checks[key] = check_summary

    artifacts: dict[str, object] = {}
    resume_ats_docx = str(result.get("resume_ats_docx") or "").strip()
    if resume_ats_docx:
        ats_docx_path = Path(resume_ats_docx)
        if ats_docx_path.exists():
            try:
                document = Document(ats_docx_path)
                paragraph_text = [
                    paragraph.text.strip()
                    for paragraph in document.paragraphs
                    if paragraph.text.strip()
                ]
                headings: list[str] = []
                seen_headings: set[str] = set()
                for text in paragraph_text:
                    if text not in ATS_STANDARD_HEADINGS or text in seen_headings:
                        continue
                    seen_headings.add(text)
                    headings.append(text)

                fonts = sorted(
                    {
                        run.font.name
                        for paragraph in document.paragraphs
                        for run in paragraph.runs
                        if run.font.name
                    }
                )
                header_text: list[str] = []
                footer_text: list[str] = []
                column_counts: list[int] = []
                for section in document.sections:
                    header_text.extend(
                        paragraph.text.strip()
                        for paragraph in section.header.paragraphs
                        if paragraph.text.strip()
                    )
                    footer_text.extend(
                        paragraph.text.strip()
                        for paragraph in section.footer.paragraphs
                        if paragraph.text.strip()
                    )
                    cols = section._sectPr.find(qn("w:cols"))
                    num_columns = cols.get(qn("w:num")) if cols is not None else None
                    column_counts.append(int(num_columns) if num_columns else 1)

                artifacts.update(
                    {
                        "headings": headings,
                        "tables": len(document.tables),
                        "inline_shapes": len(document.inline_shapes),
                        "fonts": fonts,
                        "nonempty_header_text": header_text,
                        "nonempty_footer_text": footer_text,
                        "column_counts": column_counts,
                    }
                )
            except Exception as exc:  # noqa: BLE001
                artifacts["docx_audit_error"] = str(exc)

    resume_pdf = str(result.get("resume_pdf") or "").strip()
    if resume_pdf:
        pdf_path = Path(resume_pdf)
        if pdf_path.exists():
            try:
                artifacts["pdf_page_count"] = len(PdfReader(str(pdf_path)).pages)
            except Exception as exc:  # noqa: BLE001
                artifacts["pdf_audit_error"] = str(exc)

    return {
        "generation_status": result.get("status"),
        "must_have_keyword_override_applied": bool(
            result.get("must_have_keyword_override_applied")
        ),
        "checks": checks,
        "artifacts": artifacts,
    }


def write_ats_audit_summary(package_dir: Path, ats_audit: dict[str, object]) -> Path:
    """Persist the ATS audit next to the packaged resume and cover letter artifacts."""
    output_path = package_dir / "ATS Audit Summary.json"
    output_path.write_text(json.dumps(ats_audit, indent=2), encoding="utf-8")
    return output_path


def format_ats_audit_statuses(ats_audit: object) -> str:
    """Format ATS check statuses for batch summary output."""
    if not isinstance(ats_audit, dict):
        return ""
    checks = ats_audit.get("checks")
    if not isinstance(checks, dict):
        return ""

    statuses: list[str] = []
    for key in ATS_AUDIT_CHECK_KEYS:
        raw_check = checks.get(key)
        if not isinstance(raw_check, dict):
            continue
        status = str(raw_check.get("status") or "").strip()
        if status:
            statuses.append(f"{key}={status}")
    return ", ".join(statuses)


def generate_manual_application(
    request: ManualApplicationRequest,
    context: ApplicationGenerationContext,
) -> tuple[dict[str, object], str, Path]:
    """Generate a tailored resume, cover letter, and packaged output for one role."""
    result = generate_tailored_resume_for_job(
        job_url=request.job_url,
        job_description=request.job_description,
        company=request.company,
        role=request.role,
        match_score=request.match_score,
        template_hint=request.template_hint,
        preferred_accomplishment_ids=request.preferred_accomplishment_ids,
        model_override=context.model,
        allow_must_have_keyword_override=request.allow_must_have_keyword_override,
    )
    if not result or result.get("status") != "success":
        raise RuntimeError(f"Resume generation failed: {result}")

    resume_source_path = Path(str(result["resume_source"]))
    cover_letter_text = generate_cover_letter(
        context.client,
        request.job_description,
        resume_source_path.read_text(encoding="utf-8"),
        request.company,
        request.role,
        context.model,
        guidance=CoverLetterGuidance(
            industry_focus=request.industry_focus,
            company_context=request.cover_letter_focus,
            emphasis_points=request.tailoring_notes,
        ),
    ).strip()
    if not cover_letter_text:
        raise RuntimeError("Cover letter generation returned empty text.")

    manual_review_notes: list[str] = []
    if bool(result.get("must_have_keyword_override_applied")):
        manual_review_notes.append(
            "Override applied: this package was exported for manual review even though the must-have keyword gate still flagged missing exact JD phrases."
        )
        must_have_keyword_gate = result.get("must_have_keyword_gate")
        if isinstance(must_have_keyword_gate, Mapping):
            missing_phrases = must_have_keyword_gate.get("missing_phrases")
            if isinstance(missing_phrases, list) and missing_phrases:
                manual_review_notes.append(
                    "Missing phrases at generation time: " + ", ".join(
                        str(phrase).strip()
                        for phrase in missing_phrases[:5]
                        if str(phrase).strip()
                    )
                )

    package = package_application_artifacts(
        package_root=request.package_root,
        company=request.company,
        role=request.role,
        resume_pdf=Path(str(result["resume_pdf"])) if result.get("resume_pdf") else None,
        resume_docx=Path(str(result["resume_docx"])) if result.get("resume_docx") else None,
        resume_ats_docx=(
            Path(str(result["resume_ats_docx"])) if result.get("resume_ats_docx") else None
        ),
        cover_letter_text=cover_letter_text,
        job_url=request.job_url,
        analysis_report=Path(str(result["analysis_report"])) if result.get("analysis_report") else None,
        resume_source=resume_source_path,
        include_supporting_artifacts=request.include_supporting_artifacts,
        manual_review_notes=manual_review_notes,
    )
    result["package_dir"] = str(package.package_dir)
    result["preferred_resume_docx"] = (
        str(package.preferred_resume_docx) if package.preferred_resume_docx else None
    )
    result["cover_letter_txt"] = str(package.cover_letter_txt)
    result["cover_letter_docx"] = str(package.cover_letter_docx)
    ats_audit = build_ats_audit_summary(result)
    result["ats_audit"] = ats_audit
    result["ats_audit_summary"] = str(write_ats_audit_summary(package.package_dir, ats_audit))
    return result, cover_letter_text, package.package_dir


def write_batch_apply_links(
    results: list[dict[str, object]],
    manifest: ApplicationBatchManifest,
    output_root: Path,
) -> Path:
    """Write a batch-level apply-link index for quick manual review."""
    output_root.mkdir(parents=True, exist_ok=True)
    summary_path = output_root / "apply_links.txt"
    jobs_by_key = {(job.company, job.role): job for job in manifest.jobs}
    successful_count = sum(1 for result in results if result.get("status") == "success")

    lines = [
        "ATS Sniper Batch Apply Links",
        f"Generated: {datetime.now().isoformat()}",
        f"Requested jobs: {len(manifest.jobs)}",
        f"Successful packages: {successful_count}",
        "",
    ]

    for index, result in enumerate(results, start=1):
        company = str(result.get("company", "Unknown")).strip() or "Unknown"
        role = str(result.get("role", "Unknown")).strip() or "Unknown"
        status = str(result.get("status", "unknown")).strip() or "unknown"
        package_dir = str(result.get("package_dir", "")).strip()
        error = str(result.get("error", "")).strip()
        ats_audit_statuses = format_ats_audit_statuses(result.get("ats_audit"))
        job = jobs_by_key.get((company, role))
        job_url = job.job_url if job else ""

        lines.append(f"{index}. {company} - {role}")
        lines.append(f"   Status: {status}")
        if job_url:
            lines.append(f"   URL: {job_url}")
        if package_dir:
            lines.append(f"   Package: {package_dir}")
        if ats_audit_statuses:
            lines.append(f"   ATS Audit: {ats_audit_statuses}")
        if error:
            lines.append(f"   Error: {error}")
        lines.append("")

    summary_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return summary_path


def generate_batch_applications(
    manifest_path: Path,
    *,
    package_root_override: Path | None = None,
    model_override: str | None = None,
    include_supporting_artifacts_override: bool | None = None,
    min_match_score: int | None = None,
    max_match_score: int | None = None,
) -> list[dict[str, object]]:
    """Generate application packs for every job listed in a markdown batch manifest."""
    manifest = parse_application_batch_manifest(manifest_path)
    filtered_jobs = filter_batch_jobs_by_match_score(
        manifest.jobs,
        min_match_score=min_match_score,
        max_match_score=max_match_score,
    )
    if not filtered_jobs:
        print("No batch jobs matched the requested match-score filter.")
        return []
    filtered_manifest = ApplicationBatchManifest(
        package_root=manifest.package_root,
        include_supporting_artifacts=manifest.include_supporting_artifacts,
        jobs=filtered_jobs,
    )
    if len(filtered_jobs) != len(manifest.jobs):
        print(
            f"Processing {len(filtered_jobs)} of {len(manifest.jobs)} manifest jobs "
            f"after match-score filtering."
        )
    context = build_generation_context(model_override)
    results: list[dict[str, object]] = []
    batch_output_root = package_root_override or filtered_manifest.package_root

    for job in filtered_manifest.jobs:
        print(f"\n=== {job.company} | {job.role} ===")
        try:
            request = build_request_from_batch_job(
                job,
                filtered_manifest,
                package_root_override=package_root_override,
                include_supporting_artifacts_override=include_supporting_artifacts_override,
            )
            result, _cover_letter_text, package_dir = generate_manual_application(request, context)
            results.append(
                {
                    "company": job.company,
                    "role": job.role,
                    "status": "success",
                    "model": context.model,
                    "package_dir": str(package_dir),
                    "resume_pdf": result.get("resume_pdf"),
                    "resume_docx": result.get("resume_docx"),
                    "resume_ats_docx": result.get("resume_ats_docx"),
                    "preferred_resume_docx": result.get("preferred_resume_docx"),
                    "cover_letter_txt": result.get("cover_letter_txt"),
                    "cover_letter_docx": result.get("cover_letter_docx"),
                    "resume_source": result.get("resume_source"),
                    "ats_audit": result.get("ats_audit"),
                    "ats_audit_summary": result.get("ats_audit_summary"),
                }
            )
        except Exception as exc:  # noqa: BLE001
            results.append(
                {
                    "company": job.company,
                    "role": job.role,
                    "status": "error",
                    "error": str(exc),
                }
            )
            print(f"ERROR: {exc}")

        summary_path = write_batch_apply_links(results, filtered_manifest, batch_output_root)
        print(f"Apply link summary: {summary_path}")
    return results


def generate_review_csv_applications(
    review_csv_path: Path,
    *,
    queue_ranks: tuple[int, ...] = (),
    package_root_override: Path | None = None,
    model_override: str | None = None,
    include_supporting_artifacts_override: bool | None = None,
    min_match_score: int | None = None,
    max_match_score: int | None = None,
    actionable_only: bool = True,
    limit: int | None = None,
) -> list[dict[str, object]]:
    """Generate application packs directly from the review-queue CSV export."""
    selected_rows = filter_review_csv_rows(
        load_review_csv_rows(review_csv_path),
        queue_ranks=queue_ranks,
        min_match_score=min_match_score,
        max_match_score=max_match_score,
        actionable_only=actionable_only,
        limit=limit,
    )
    if not selected_rows:
        print("No review CSV rows matched the requested filters.")
        return []

    include_supporting_artifacts = bool(include_supporting_artifacts_override)
    batch_output_root = package_root_override or DEFAULT_REVIEW_PACKAGE_ROOT
    manifest = build_manifest_from_review_csv_rows(
        selected_rows,
        package_root=batch_output_root,
        include_supporting_artifacts=include_supporting_artifacts,
    )
    context = build_generation_context(model_override)
    state = load_state()
    results: list[dict[str, object]] = []

    for row in selected_rows:
        company = str(row.get("Company", "Unknown")).strip() or "Unknown"
        role = str(row.get("Title", "Unknown")).strip() or "Unknown"
        print(f"\n=== {company} | {role} ===")
        try:
            request = build_request_from_review_csv_row(
                row,
                state=state,
                package_root=batch_output_root,
                include_supporting_artifacts=include_supporting_artifacts,
            )
            result, _cover_letter_text, package_dir = generate_manual_application(request, context)
            results.append(
                {
                    "company": company,
                    "role": role,
                    "status": "success",
                    "model": context.model,
                    "package_dir": str(package_dir),
                    "resume_pdf": result.get("resume_pdf"),
                    "resume_docx": result.get("resume_docx"),
                    "resume_ats_docx": result.get("resume_ats_docx"),
                    "preferred_resume_docx": result.get("preferred_resume_docx"),
                    "cover_letter_txt": result.get("cover_letter_txt"),
                    "cover_letter_docx": result.get("cover_letter_docx"),
                    "resume_source": result.get("resume_source"),
                    "ats_audit": result.get("ats_audit"),
                    "ats_audit_summary": result.get("ats_audit_summary"),
                }
            )
        except Exception as exc:  # noqa: BLE001
            results.append(
                {
                    "company": company,
                    "role": role,
                    "status": "error",
                    "error": str(exc),
                }
            )
            print(f"ERROR: {exc}")

    summary_path = write_batch_apply_links(results, manifest, batch_output_root)
    print(f"Apply link summary: {summary_path}")
    return results


def main() -> int:
    """Run the manual application package flow from the command line."""
    args = parse_arguments()

    try:
        if args.batch_file:
            results = generate_batch_applications(
                Path(args.batch_file),
                package_root_override=resolve_package_root_override(args.package_root),
                model_override=args.model,
                include_supporting_artifacts_override=(
                    True if args.include_supporting_artifacts else None
                ),
                min_match_score=args.min_match_score,
                max_match_score=args.max_match_score,
            )
            print("\n=== BATCH RESULTS ===")
            for result in results:
                print(result)
            return 0 if all(result.get("status") == "success" for result in results) else 1

        context = build_generation_context(args.model)
        request = build_request(args)
        result, _cover_letter_text, package_dir = generate_manual_application(request, context)
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}")
        return 1

    print(f"Application pack created: {package_dir}")
    print(f"Preferred ATS Upload: {result.get('preferred_resume_docx')}")
    print(f"Resume PDF: {result.get('resume_pdf')}")
    print(f"Resume DOCX: {result.get('resume_docx')}")
    print(f"Resume ATS DOCX: {result.get('resume_ats_docx')}")
    print(f"Resume source: {result.get('resume_source')}")
    print(f"Analysis report: {result.get('analysis_report')}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

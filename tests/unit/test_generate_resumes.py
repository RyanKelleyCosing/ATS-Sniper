"""Unit tests for resume rendering layout helpers."""

from pathlib import Path
import sys

from docx import Document
from reportlab.pdfbase.pdfmetrics import stringWidth


ATS_ROOT = Path(__file__).resolve().parents[2]
RESUME_ROOT = ATS_ROOT.parent

for candidate in (ATS_ROOT, RESUME_ROOT):
    if str(candidate) not in sys.path:
        sys.path.insert(0, str(candidate))

from generate_resumes import (
    audit_resume_layouts,
    calculate_pdf_layout_metrics,
    condense_resume_data_for_one_page,
    compress_left_label,
    create_ats_docx_resume,
    create_docx_resume,
    fit_text_to_single_line,
    get_contact_entries,
    parse_accomplishments,
    parse_resume_source,
    wrap_text_to_width,
)
from params.experience_dates import RESURGENT_PROMOTION_BULLET_SOURCE


def test_parse_resume_source_normalizes_known_experience_dates(tmp_path: Path) -> None:
    resume_source = tmp_path / "resume_dates.md"
    resume_source.write_text(
        """---
name: Test Candidate
role: DevOps Engineer
contact: {}
skills: {}
certifications: []
education: Test University
summary: Test summary
---

## Experience

### Resurgent Capital Services | Software Support Analyst II | Aug 2021 - Mar 2026
- Example bullet. Technologies: Azure DevOps, Kubernetes

### Silco | Database Analyst (Contract) | Aug 2020 - May 2021
- Example bullet. Technologies: SQL Server, SSIS

### RIBBIT.AI | Database Analyst (Contract) | Feb 2020 - Aug 2020
- Example bullet. Technologies: SQL Server, ETL
""",
        encoding="utf-8",
    )

    parsed = parse_resume_source(resume_source, accomplishments={})

    assert parsed["experience"][0]["title"] == "Software Support Analyst II"
    assert parsed["experience"][0]["dates"] == "Aug 2021 - Mar 2026"
    assert parsed["experience"][1]["dates"] == "Jan 2021 - May 2021"
    assert parsed["experience"][2]["dates"] == "Aug 2020 - Jan 2021"


def test_parse_resume_source_accepts_work_experience_heading(tmp_path: Path) -> None:
    resume_source = tmp_path / "resume_work_experience.md"
    resume_source.write_text(
        """---
name: Test Candidate
role: DevOps Engineer
contact: {}
skills: {}
certifications: []
education: Test University
summary: Test summary
---

## Work Experience

### Example Co | Engineer | 2024 - Present
- Automated deployments. Technologies: Azure DevOps, Bicep
""",
        encoding="utf-8",
    )

    parsed = parse_resume_source(resume_source, accomplishments={})

    assert parsed["experience"][0]["company"] == "Example Co"
    assert parsed["experience"][0]["bullets"] == [("Automated deployments", "Azure DevOps, Bicep")]


def test_wrap_text_to_width_breaks_long_urls() -> None:
    max_width = (2.0 * 72) - (2 * 0.3 * 72)

    wrapped_lines = wrap_text_to_width(
        "github.com/example-very-long-handle-for-wrap-test",
        "Helvetica",
        11,
        max_width,
    )

    assert len(wrapped_lines) > 1
    assert all(stringWidth(line, "Helvetica", 11) <= max_width for line in wrapped_lines)


def test_audit_resume_layouts_reports_compact_resume_as_ok(tmp_path: Path) -> None:
    accomplishments = parse_accomplishments(RESUME_ROOT / "accomplishments.md")
    resume_source = tmp_path / "resume_layout_audit.md"
    resume_source.write_text(
    """---
name: Test Candidate
role: DevOps Engineer
contact:
    email: test@example.com
    location: Cincinnati, OH
skills:
    Cloud Architecture:
        - Azure
        - Kubernetes
        - Docker
    CI/CD Automation:
        - Azure DevOps
        - GitHub Actions
    Operations:
        - Azure Monitor
        - Grafana
certifications: []
education: Test University
summary: DevOps engineer with cloud automation and monitoring experience.
technical_environment: Azure, Kubernetes, Docker, Azure DevOps, Azure Monitor
---

## Experience

### Example Co | Engineer | 2024 - Present
- Automated Azure deployments and monitoring for production services. Technologies: Azure DevOps, Azure Monitor, Kubernetes
- Reduced environment drift with infrastructure automation and repeatable release workflows. Technologies: Docker, GitHub Actions, Azure
""",
        encoding="utf-8",
    )

    audit_results = audit_resume_layouts([resume_source], accomplishments)

    assert len(audit_results) == 1
    assert audit_results[0]["status"] == "ok"
    assert audit_results[0]["estimated_pages"] == 1


def test_compress_left_label_prefers_whole_words() -> None:
    assert compress_left_label("MONITORING & OPS") == "MONITORING"
    assert compress_left_label("CLOUD & INFRASTRUCTURE") == "CLOUD PLATFORM"
    assert compress_left_label("SCRIPTING & DATABASES") == "SCRIPTING & DATA"
    assert compress_left_label("DATA MANAGEMENT") == "DATA"


def test_get_contact_entries_shortens_profile_links_and_mailto_email() -> None:
    contact = {
        "phone": "(760) 216-7729",
        "email": "candidate@example.com",
        "location": "Cincinnati, OH",
        "linkedin": "https://www.linkedin.com/in/ryan-kelley-it/",
        "github": "https://github.com/example",
    }

    entries = get_contact_entries(contact)

    assert entries[0] == ("(760) 216-7729", None)
    assert entries[1] == ("candidate@example.com", "mailto:candidate@example.com")
    assert entries[3] == ("linkedin/ryan-kelley-it", "https://www.linkedin.com/in/ryan-kelley-it/")
    assert entries[4] == ("github/example", "https://github.com/example")


def test_fit_text_to_single_line_keeps_platform_labels_unbroken() -> None:
    max_width = (2.0 * 72) - (2 * 0.24 * 72)

    label, font_size = fit_text_to_single_line(
        "CLOUD ARCHITECTURE",
        "Helvetica-Bold",
        10.5,
        max_width,
    )

    assert label == "CLOUD ARCHITECTURE"
    assert stringWidth(label, "Helvetica-Bold", font_size) <= max_width


def test_parse_resume_source_reads_selected_achievements(tmp_path: Path) -> None:
    resume_source = tmp_path / "resume_test.md"
    resume_source.write_text(
        """---
name: Test Candidate
role: Cloud Engineer II
contact: {}
skills: {}
certifications: []
education: Test University
summary: Test summary
technical_environment: Azure, Bicep, AKS, Terraform, PowerShell
selected_achievements:
  - bullet: Reduced Azure spend by 22% through rightsizing
    technologies: Azure Cost Management, PowerShell
---

## Experience

### Example Co | Engineer | 2024 - Present
- Reduced Azure spend by 22% through rightsizing. Technologies: Azure Cost Management, PowerShell
""",
        encoding="utf-8",
    )

    parsed = parse_resume_source(resume_source, accomplishments={})

    assert parsed["selected_achievements"] == [
        ("Reduced Azure spend by 22% through rightsizing", "Azure Cost Management, PowerShell")
    ]
    assert parsed["technical_environment"] == "Azure, Bicep, AKS, Terraform, PowerShell"


def test_calculate_pdf_layout_metrics_uses_selected_achievements_to_fill_space() -> None:
    data = {
        "name": "TEST CANDIDATE",
        "title": "Cloud Engineer II",
        "contact": {"email": "test@example.com"},
        "skills": {"Cloud Architecture": ["Azure", "Bicep", "AKS"]},
        "certifications": [],
        "education": "Test University",
        "summary": "Cloud engineer with strong Azure and automation experience.",
        "selected_achievements": [
            ("Reduced Azure spend by 22% through rightsizing", "Azure Cost Management, PowerShell"),
        ],
        "technical_environment": "Azure, Bicep, AKS, Terraform, PowerShell, Azure Monitor",
        "experience": [
            {
                "company": "Example Co",
                "title": "Engineer",
                "dates": "2024 - Present",
                "bullets": [("Automated Azure deployments", "Azure DevOps, Bicep")],
            }
        ],
    }

    metrics = calculate_pdf_layout_metrics(data)

    assert metrics.selected_achievements == [
        ("Reduced Azure spend by 22% through rightsizing", "Azure Cost Management, PowerShell")
    ]
    assert metrics.technical_environment.startswith("Azure")
    assert metrics.white_space >= 0


def test_condense_resume_data_for_one_page_trims_overflow() -> None:
    long_bullet = (
        "Architected a multi-stage platform automation workflow across hybrid environments, "
        "coordinating monitoring, deployment validation, rollback controls, and incident recovery "
        "for business-critical services while reducing manual effort and improving reliability"
    )
    data = {
        "name": "TEST CANDIDATE",
        "title": "Site Reliability Engineer II",
        "contact": {"email": "test@example.com"},
        "skills": {"Reliability Ops": ["Azure Monitor", "Grafana", "PowerShell"]},
        "certifications": [],
        "education": "Test University",
        "summary": (
            "Site reliability engineer with deep observability and automation experience across hybrid platforms. "
            "Leads incident response improvements, deployment hardening, and cost-aware platform operations."
        ),
        "selected_achievements": [],
        "technical_environment": "Azure Monitor, Grafana, PowerShell, Kubernetes, Docker",
        "experience": [
            {
                "company": "Example Co",
                "title": "Senior Software Support Analyst II",
                "dates": "2021 - Present",
                "bullets": [(long_bullet, "Azure Monitor, Grafana, PowerShell") for _ in range(9)],
            },
            {
                "company": "Silco",
                "title": "Database Analyst",
                "dates": "2020 - 2021",
                "bullets": [(long_bullet, "SQL Server, SSIS, T-SQL") for _ in range(3)],
            },
            {
                "company": "RIBBIT.AI",
                "title": "Database Analyst",
                "dates": "2020",
                "bullets": [(long_bullet, "SQL Server, SSRS, ETL") for _ in range(3)],
            },
        ],
    }

    initial_metrics = calculate_pdf_layout_metrics(data)
    condensed = condense_resume_data_for_one_page(data)
    condensed_metrics = calculate_pdf_layout_metrics(condensed)

    assert initial_metrics.estimated_pages > 1
    assert condensed_metrics.estimated_pages == 1


def test_create_ats_docx_resume_uses_single_column_sections(tmp_path: Path) -> None:
    output_path = tmp_path / 'resume_ats.docx'
    data = {
        'name': 'TEST CANDIDATE',
        'title': 'IT Platform Engineer',
        'contact': {
            'email': 'test@example.com',
            'location': 'Cincinnati, OH',
            'linkedin': 'https://www.linkedin.com/in/test-candidate/',
        },
        'skills': {
            'Cloud Architecture': ['Azure', 'AWS', 'Bicep'],
            'Operations': ['Windows Server', 'Azure Monitor'],
        },
        'certifications': ['Azure Administrator Associate'],
        'education': 'Test University',
        'summary': 'Platform engineer with Azure, automation, and operational support experience.',
        'technical_environment': 'Azure, AWS, Windows Server, Active Directory, Azure Monitor',
        'experience': [
            {
                'company': 'Example Co',
                'title': 'Software Support Analyst II',
                'dates': '2021 - Present',
                'bullets': [
                    ('Administered Azure platform services for business-critical applications.', 'Azure, Windows Server'),
                ],
            }
        ],
    }

    create_ats_docx_resume('IT_Platform_Engineer', data, output_path)

    document = Document(output_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    assert output_path.exists()
    assert document.tables == []
    assert paragraphs[0] == 'TEST CANDIDATE'
    assert any('Cincinnati, OH' in paragraph for paragraph in paragraphs)
    assert any('www.linkedin.com/in/test-candidate' in paragraph for paragraph in paragraphs)
    assert 'PROFESSIONAL SUMMARY' in paragraphs
    assert 'TECHNICAL SKILLS' in paragraphs
    assert 'WORK EXPERIENCE' in paragraphs
    assert 'EDUCATION' in paragraphs
    assert 'CERTIFICATIONS' in paragraphs
    assert any(paragraph.startswith('Cloud Architecture: Azure, AWS, Bicep') for paragraph in paragraphs)
    assert any(paragraph.startswith('Additional Tools: Azure, AWS') for paragraph in paragraphs)
    assert 'Azure Administrator Associate' in paragraphs
    assert 'Example Co | 2021 - Present' in paragraphs
    assert 'Software Support Analyst II' in paragraphs
    assert any(paragraph.startswith('- Administered Azure platform services') for paragraph in paragraphs)
    assert not any(' @ ' in paragraph for paragraph in paragraphs)

    explicit_body_fonts = {
        run.font.name
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text.strip() and run.font.name
    }
    assert explicit_body_fonts == {'Calibri'}


def test_create_ats_docx_resume_simplifies_resurgent_promotion_timeline(tmp_path: Path) -> None:
    output_path = tmp_path / 'resume_ats.docx'
    data = {
        'name': 'TEST CANDIDATE',
        'title': 'Cloud Engineer',
        'contact': {'email': 'test@example.com'},
        'skills': {'Cloud': ['Azure', 'Bicep']},
        'certifications': [],
        'education': 'Test University',
        'summary': 'Cloud engineer with Azure automation experience.',
        'technical_environment': 'Azure, Bicep',
        'experience': [
            {
                'company': 'Resurgent Capital Services',
                'title': 'Software Support Analyst II',
                'dates': 'Aug 2021 - Mar 2026',
                'bullets': [
                    (RESURGENT_PROMOTION_BULLET_SOURCE, 'Azure, APIs'),
                ],
            }
        ],
    }

    create_ats_docx_resume('Cloud_Engineer', data, output_path)

    document = Document(output_path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    assert any(
        paragraph.startswith(
            '- Earned two promotions while expanding scope into'
        )
        for paragraph in paragraphs
    )
    assert not any('IT Analyst (Aug 2021 - Sep 2022)' in paragraph for paragraph in paragraphs)


def test_create_docx_resume_uses_separate_right_column_paragraphs_and_spacing(tmp_path: Path) -> None:
    output_path = tmp_path / 'resume_styled.docx'
    data = {
        'name': 'TEST CANDIDATE',
        'title': 'Implementation Engineer',
        'contact': {
            'email': 'test@example.com',
            'location': 'Cincinnati, OH',
        },
        'skills': {
            'Implementation Delivery': ['Deployment Validation', 'Release Coordination'],
            'Support Operations': ['Troubleshooting', 'Incident Triage'],
        },
        'certifications': ['Azure Developer Associate'],
        'education': 'Test University',
        'summary': 'Implementation engineer with API integration, support, and deployment experience.',
        'selected_achievements': [],
        'technical_environment': 'Azure, SQL Server, REST APIs',
        'experience': [
            {
                'company': 'Example Co',
                'title': 'Software Support Analyst II',
                'dates': '2021 - Present',
                'bullets': [
                    ('Resolved API integration failures and reduced onboarding delays by 45%', 'REST APIs, SQL Server'),
                ],
            }
        ],
    }

    create_docx_resume('Implementation_Support_Engineer', data, output_path)

    document = Document(output_path)
    right_cell = document.tables[0].rows[0].cells[1]
    right_paragraphs = [paragraph for paragraph in right_cell.paragraphs if paragraph.text.strip()]
    right_texts = [paragraph.text for paragraph in right_paragraphs]

    assert output_path.exists()
    assert right_texts[:4] == [
        'TEST CANDIDATE',
        'Implementation Engineer',
        'PROFESSIONAL SUMMARY',
        'Implementation engineer with API integration, support, and deployment experience.',
    ]
    assert 'WORK EXPERIENCE' in right_texts
    work_experience_paragraph = next(
        paragraph for paragraph in right_paragraphs if paragraph.text == 'WORK EXPERIENCE'
    )
    assert work_experience_paragraph.paragraph_format.space_after.pt > 0